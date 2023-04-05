import os
import re
import docx
import time
import json
import pysrt
import requests
import numpy as np
import streamlit as st
import seaborn as sns
import matplotlib.pyplot as plt
from collections import defaultdict
from wikitextparser import parse, remove_markup


def styled_print(text, header=False, sleep_time=0):
    """Custom Print Function"""
    class style:
        BOLD = '\033[1m'
        UNDERLINE = '\033[4m'
        END = '\033[0m'

    if header:
        print(f'{style.BOLD}› {style.UNDERLINE}{text}{style.END}')
    else:
        print(f'    › {text}')
    time.sleep(sleep_time)


def create_dir(root_dir, new_dir, header=True):
    styled_print(
        f'creating directory ... {os.path.join(root_dir, new_dir)}', header=header)
    os.makedirs(os.path.join(root_dir, new_dir), exist_ok=True)
    return os.path.join(root_dir, new_dir)


def extract_paragraphs(file_path, min_char_count=1):
    styled_print(f"Extracting Paragraphs from {file_path}", header=False)
    paragraphs = {}
    document = docx.Document(file_path)
    for i in range(2, len(document.paragraphs)):
        if min_char_count is not None:
            if len(document.paragraphs[i].text) >= min_char_count:
                paragraphs[i] = document.paragraphs[i].text
        else:
            paragraphs[i] = document.paragraphs[i].text
    return paragraphs


def extract_text_from_srt(file_path):
    subtitle_dict = defaultdict(list)
    subtitles = pysrt.open(file_path)
    for sub in subtitles:
        subtitle_dict["start_hours"].append(sub.start.hours)
        subtitle_dict["start_minutes"].append(sub.start.minutes)
        subtitle_dict["start_seconds"].append(sub.start.seconds)
        subtitle_dict["end_hours"].append(sub.end.hours)
        subtitle_dict["end_minutes"].append(sub.end.minutes)
        subtitle_dict["end_seconds"].append(sub.end.seconds)
        subtitle_dict["text"].append(sub.text)
    return subtitle_dict


def extract_fandom_wikis(url, titles):
    def extract_sections(wikitext):
        parsed_wikitext = parse(wikitext)
        sections = {}
        for section in parsed_wikitext.sections:
            if section.title and section.title.strip():
                title = str(section.title.strip())
                title = remove_file_links(title)
                title = remove_newlines(title)
                title = remove_backslash(title)
                title = remove_markup(title)
            else:
                title = "No Title"
            if section.contents and section.contents.strip():
                contents = str(section.contents.strip())
                contents = remove_file_links(contents)
                contents = remove_newlines(contents)
                contents = remove_backslash(contents)
                contents = remove_markup(contents)
            else:
                contents = ""
            sections[title] = contents
        return sections

    data_dict = defaultdict(dict)
    response = requests.get(url + "?action=raw")
    wikitext_content = response.text
    parsed_wikitext = parse(wikitext_content)
    for section in parsed_wikitext.sections:
        if section.title in titles:
            data_dict[section.title] = extract_sections(section.contents)

    sample_dict = defaultdict(list)
    for key, value in data_dict.items():
        for k, v in value.items():
            if not v:
                continue
            sample_dict["header"].append(key)
            sample_dict["prompt"].append(k)
            sample_dict["paragraphs"].append(v)
    return sample_dict


def remove_file_links(text):
    pattern = r'\[\[File:(.*?)\]\]'
    return re.sub(pattern, '', text)


def remove_newlines(text):
    return re.sub(r"\n", "", text)


def remove_backslash(text):
    return re.sub(r"\\'", "'", text)


def save_uploadedfile(uploadedfile, path="tempDir"):
    with open(os.path.join(path, uploadedfile.name), "wb") as f:
        f.write(uploadedfile.getbuffer())
    return os.path.join(path, uploadedfile.name)


def add_to_session_state(session_state, key, value):
    if key not in session_state:
        session_state[key] = value
    return session_state


def fetch_session_state(session_state, key):
    if key in session_state:
        return session_state[key]
    else:
        return None


def plot_box_plot_hist_plot(df, column, title="Distribution Plot", figsize=(16, 16),
                            dpi=300, save_flag=False, file_path=None):
    fig, (ax_box, ax_hist) = plt.subplots(
        nrows=2,
        sharex=True,
        figsize=figsize,
        gridspec_kw={"height_ratios": (.20, .80)},
        dpi=dpi,
        constrained_layout=False
    )
    sns.boxplot(data=df, x=column, ax=ax_box)
    sns.histplot(data=df, x=column, ax=ax_hist, kde=True, bins='sqrt')
    ax_box.set(xlabel='')
    ax_box.set_facecolor('white')
    ax_hist.set_facecolor('white')
    plt.title(title)
    if save_flag:
        fig.savefig(file_path, dpi=dpi, facecolor='white')
        plt.close()
    return fig


def plot_count_plot(df, column, hue=None, title="Count Plot", figsize=(24, 24), dpi=300,
                    save_flag=False, file_path=None):
    fig, axs = plt.subplots(1, 1, figsize=figsize,
                            dpi=dpi, constrained_layout=False)
    pt = sns.countplot(data=df, x=column, hue=hue,
                       palette=sns.color_palette("Set2"))
    pt.set_xticklabels(pt.get_xticklabels(), rotation=30)
    if hue is not None:
        axs.legend(loc="upper right", title=hue)
    axs.set_facecolor('white')
    plt.title(title)
    if save_flag:
        fig.savefig(file_path, dpi=dpi, facecolor='white')
        plt.close()
    return fig
